"""
Generates a detailed Word document explaining the Adapta project for an interview.
Run: pip install python-docx && python generate_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

doc = Document()

# ──────────────────────────────────────────────────────────────────────────────
# STYLES
# ──────────────────────────────────────────────────────────────────────────────
def set_page_margins(doc, top=1, bottom=1, left=1.2, right=1.2):
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)

def style_run(run, size=11, bold=False, italic=False, color=None, font="Calibri"):
    run.font.name  = font
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(0,0,0)):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(*color)
    return para

def add_para(doc, text="", bold_prefix=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        style_run(r, bold=True, color=(63,63,63))
    r = p.add_run(text)
    style_run(r)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    r = p.add_run(text)
    style_run(r)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    style_run(r, font="Courier New", size=9, color=(20,20,80))
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    hrow = table.rows[0].cells
    for i, h in enumerate(headers):
        hrow[i].text = h
        for r in hrow[i].paragraphs[0].runs:
            r.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()

def divider(doc):
    doc.add_paragraph("─" * 80)

set_page_margins(doc)

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\nADAPTA\n")
style_run(r, size=36, bold=True, color=(79,70,229))      # indigo

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("AI-Powered Task Prioritization Platform\n")
style_run(r2, size=18, color=(100,100,100))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Complete Technical Documentation — Interview Reference\n\n")
style_run(r3, size=12, italic=True, color=(120,120,120))

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Stack: React · Java Spring Boot 3 · MongoDB · Firebase Auth · Google Gemini AI")
style_run(r4, size=11, bold=True, color=(50,50,50))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Project Overview", 1, (79,70,229))
add_para(doc,
    "Adapta is a full-stack, AI-driven task prioritization platform designed to help software "
    "engineering teams focus on the work that delivers the most value. Rather than manually "
    "ordering backlogs, Adapta automatically computes a 0–100 priority score for every task "
    "based on business impact, implementation effort, and project risk. An AI urgency layer "
    "(powered by Google Gemini) further boosts scores for tasks with nearby deadlines or "
    "heavily loaded projects. A dedicated Team Formation module uses AI to match the right "
    "people to the right tasks based on their skill sets, availability, and current workload.")

add_heading(doc, "1.1 Core Problem Being Solved", 2)
add_bullet(doc, "Engineers waste hours deciding what to work on next — Adapta makes that decision automatically.")
add_bullet(doc, "Managers lack visibility into team bandwidth — Adapta tracks workload scores per person in real time.")
add_bullet(doc, "Ad-hoc team formation leads to skill mismatches — Adapta's AI scorer ranks the best candidates for any project.")

add_heading(doc, "1.2 Key Features", 2)
add_bullet(doc, "Automatic priority scoring: (impact × 8) + (risk × 4) − (effort × 2) + urgencyBonus")
add_bullet(doc, "Deadline-aware urgency: tasks due within 3 days get +30 pts; 7 days +20 pts; 14 days +10 pts")
add_bullet(doc, "AI task description generation via Google Gemini 1.5 Flash")
add_bullet(doc, "AI-powered team formation scoring with semantic skill matching")
add_bullet(doc, "Real-time analytics: completion rates, priority distribution, weekly burndown")
add_bullet(doc, "Kanban board view (Todo / In Progress / Done) per project")
add_bullet(doc, "Multi-step onboarding wizard capturing detailed user profiles for team matching")

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "2. System Architecture", 1, (79,70,229))

add_para(doc,
    "Adapta follows a classic three-tier architecture: a React SPA (Single Page Application) "
    "in the presentation tier, a Java Spring Boot REST API in the application tier, and MongoDB "
    "as the data tier. Firebase Authentication sits as a cross-cutting concern, issuing JWT "
    "tokens on the client and validating them on the server. Google Gemini is called "
    "server-side only, keeping the API key secure.")

add_heading(doc, "2.1 High-Level Architecture Diagram", 2)
for line in [
    "┌─────────────────────────────────────────────────────────────────────┐",
    "│                         CLIENT (Browser)                            │",
    "│  React SPA (port 3000)                                              │",
    "│  ┌─────────────┐  ┌───────────────┐  ┌────────────────────────┐   │",
    "│  │ Firebase SDK│  │ Axios Instance │  │ React Router (routes)  │   │",
    "│  │  (Auth/JWT) │  │ + JWT header  │  │ / /projects /analytics │   │",
    "│  └──────┬──────┘  └───────┬───────┘  └────────────────────────┘   │",
    "└─────────┼─────────────────┼───────────────────────────────────────-┘",
    "          │  Google         │  HTTP REST (JSON)                        ",
    "          │  Firebase       │  Bearer <JWT>                            ",
    "          ▼                 ▼                                          ",
    "┌─────────────────────────────────────────────────────────────────────┐",
    "│           Spring Boot 3 REST API (port 8080)                        │",
    "│  ┌──────────────┐ ┌───────────────┐ ┌────────────────────────────┐ │",
    "│  │  CORS Filter │ │Firebase Auth  │ │   Spring Security Chain    │ │",
    "│  │              │ │  JWT Filter   │ │     (Stateless / JWT)      │ │",
    "│  └──────────────┘ └───────┬───────┘ └────────────────────────────┘ │",
    "│                           │ FirebasePrincipal(uid, email)           │",
    "│  ┌────────────────────────▼──────────────────────────────────────┐ │",
    "│  │              Controller Layer (REST)                           │ │",
    "│  │  AuthController  ProjectController  TaskController            │ │",
    "│  │  TeamController  AnalyticsController                          │ │",
    "│  └────────────────────────┬──────────────────────────────────────┘ │",
    "│                           │                                         │",
    "│  ┌────────────────────────▼──────────────────────────────────────┐ │",
    "│  │                   Service Layer                                │ │",
    "│  │  AuthService  ProjectService  TaskService (priority algo)     │ │",
    "│  │  TeamService  AnalyticsService  AIService (Gemini API)        │ │",
    "│  └────────────────────────┬──────────────────────────────────────┘ │",
    "│                           │                                         │",
    "│  ┌────────────────────────▼──────────────────────────────────────┐ │",
    "│  │              Repository Layer (Spring Data MongoDB)            │ │",
    "│  │  UserRepository  ProjectRepository  TaskRepository            │ │",
    "│  └────────────────────────┬──────────────────────────────────────┘ │",
    "└───────────────────────────┼─────────────────────────────────────────┘",
    "                            │ MongoDB Wire Protocol                     ",
    "                            ▼                                           ",
    "┌─────────────────────────────────────────────────────────────────────┐",
    "│                   MongoDB Database: adapta                          │",
    "│         Collections: users │ projects │ tasks                      │",
    "└─────────────────────────────────────────────────────────────────────┘",
    "                                          ▲                            ",
    "                AIService ───────────────►│ Google Gemini 1.5 Flash   ",
    "                        (HTTPS REST API)  │ generativelanguage API    ",
]:
    add_code(doc, line)

# ══════════════════════════════════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "3. Technology Stack", 1, (79,70,229))

add_table(doc,
    ["Layer", "Technology", "Version", "Why Chosen"],
    [
        ["Frontend Framework",  "React",               "18.x",  "Component-based UI, large ecosystem, hooks for state management"],
        ["Frontend Routing",    "React Router DOM",    "6.x",   "Declarative client-side routing with nested routes and params"],
        ["HTTP Client",         "Axios",               "1.x",   "Interceptor support for auto-attaching Firebase JWT tokens"],
        ["CSS Methodology",     "Vanilla CSS + Custom Properties", "-", "Full control over design system; no build overhead like Tailwind"],
        ["Auth (Client)",       "Firebase Auth SDK",   "10.x",  "Google OAuth out-of-box, handles token refresh automatically"],
        ["Backend Framework",   "Spring Boot",         "3.2.4", "Production-grade REST, auto-config, mature ecosystem"],
        ["Language",            "Java",                "17",    "LTS release, records/sealed classes, strong type safety"],
        ["Auth (Server)",       "Firebase Admin SDK",  "9.2.0", "Server-side JWT verification without storing passwords"],
        ["Database",            "MongoDB",             "7.x",   "Flexible documents fit evolving schema; no JOIN overhead for task queries"],
        ["ODM (ORM equiv.)",    "Spring Data MongoDB", "4.x",   "Repository pattern, derived query methods, zero boilerplate CRUD"],
        ["AI Provider",         "Google Gemini 1.5 Flash", "-", "Fast inference, large context window, free tier available"],
        ["Security",            "Spring Security",     "6.x",   "Filter chain integration for JWT, CORS, CSRF config"],
        ["Font",                "Inter (Google Fonts)", "-",    "Clean, highly readable variable font — industry standard for SaaS"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. FRONTEND ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "4. Frontend Architecture (React)", 1, (79,70,229))

add_heading(doc, "4.1 Project Structure", 2)
for line in [
    "frontend/src/",
    "  ├── index.js          — App entry; wraps <App> in <BrowserRouter>",
    "  ├── index.css         — Global design system (CSS custom properties)",
    "  ├── App.js            — Auth-aware routing, Layout wrapper",
    "  ├── services/",
    "  │   ├── api.js        — Axios instance with Firebase JWT interceptor",
    "  │   └── firebase.js   — Firebase SDK initialization",
    "  ├── components/",
    "  │   ├── layout/",
    "  │   │   ├── Sidebar.js/css    — Fixed dark sidebar with icon nav",
    "  │   │   └── Layout.js/css     — Shell: Sidebar + sticky Topbar",
    "  │   └── ui/",
    "  │       ├── TaskCard.js/css   — Priority-colored task card",
    "  │       ├── StatCard.js/css   — KPI metric card with trend indicator",
    "  │       ├── PriorityBadge.js  — Color-coded priority badge",
    "  │       ├── ProgressRing.js   — SVG donut chart",
    "  │       ├── Modal.js/css      — Accessible overlay modal",
    "  │       └── EmptyState.js     — Empty state placeholder",
    "  └── pages/",
    "      ├── Login.js            — Split-screen auth page",
    "      ├── Onboarding.js       — 3-step profile wizard",
    "      ├── Dashboard.js        — Stat cards + task list + rings",
    "      ├── Projects.js         — Project card grid + create modal",
    "      ├── ProjectDetail.js    — Kanban board per project",
    "      ├── Analytics.js        — Charts: bars, rings, timeline",
    "      ├── TeamFormation.js    — AI team builder",
    "      └── styles/             — Co-located CSS per page",
]:
    add_code(doc, line)

add_heading(doc, "4.2 Design System (index.css)", 2)
add_para(doc,
    "The entire UI is built on a CSS custom properties (CSS variables) design system defined "
    "in index.css. This means every component references tokens like var(--color-primary), "
    "var(--space-4), var(--radius-lg) rather than hard-coded values. This makes global "
    "theming changes a single-line update.")

add_bullet(doc, "Color palette: Indigo (#6366f1) primary, Violet (#8b5cf6) accent, dark backgrounds (#0f172a, #1e293b)")
add_bullet(doc, "Typography: Inter font (Google Fonts) at weights 400/500/600/700; font sizes xs through 2xl")
add_bullet(doc, "Spacing scale: --space-1 (4px) through --space-10 (40px) — 4px base unit")
add_bullet(doc, "Shadows: sm/md/lg/xl variants tuned for dark backgrounds")
add_bullet(doc, "Border radii: sm(4px) / md(6px) / lg(10px) / xl(16px) / full(9999px)")
add_bullet(doc, "Transitions: fast (150ms ease) / base (200ms ease) used consistently")
add_bullet(doc, "Reusable classes: .btn-primary, .btn-secondary, .form-control, .card, .badge, .chip, .spinner")

add_heading(doc, "4.3 Authentication Flow (Frontend)", 2)
add_para(doc,
    "Firebase Auth SDK handles all authentication on the client. The flow is:")
for step in [
    "1. User clicks 'Continue with Google' on Login.js → Firebase Google popup opens",
    "2. Firebase returns a user object with a short-lived ID Token (JWT)",
    "3. Axios interceptor (api.js) calls user.getIdToken() before EVERY request and attaches it as Authorization: Bearer <token>",
    "4. Firebase auto-refreshes the token every 60 minutes — the interceptor always gets a fresh token",
    "5. App.js listens to auth.onAuthStateChanged() — if no user, redirects to /login",
    "6. After login, GET /auth/me is called — if hasProfile=false, redirects to /onboarding",
]:
    add_bullet(doc, step)

add_heading(doc, "4.4 Routing (App.js)", 2)
add_para(doc,
    "App.js uses React Router v6 with a two-tier routing strategy: bare routes (no sidebar/topbar) "
    "for Login and Onboarding, and Layout-wrapped routes for all authenticated pages.")
add_code(doc, "Bare routes:  /login  →  <Login />")
add_code(doc, "              /onboarding  →  <Onboarding />")
add_code(doc, "Layout routes: /  →  <Dashboard />")
add_code(doc, "               /projects  →  <Projects />")
add_code(doc, "               /projects/:id  →  <ProjectDetail />  (dynamic segment)")
add_code(doc, "               /analytics  →  <Analytics />")
add_code(doc, "               /team  →  <TeamFormation />")

add_heading(doc, "4.5 Pages — Detailed Explanation", 2)

pages = [
    ("Login.js",
     "Split-screen layout. Left panel: dark indigo gradient with Adapta brand, tagline, and 3 feature "
     "highlight rows (Smart Prioritization, Real-time Analytics, Team Formation). Right panel: white "
     "card with 'Welcome back' heading and a Google Sign-In button. Firebase signInWithPopup() handles "
     "the OAuth flow. A loading spinner appears while authentication is in progress. On success, App.js "
     "auth listener fires and navigates to dashboard or onboarding."),
    ("Onboarding.js",
     "3-step wizard with an animated step indicator (dots + progress bar). Step 1: Personal info "
     "(name, title, department, location, timezone). Step 2: Skills & experience (years, skills "
     "as tags, certifications). Step 3: Availability & preferences (hours/week, work preferences). "
     "On submit, POST /profile/complete is called with the full profile object. This data is critical "
     "for the AI team formation scoring."),
    ("Dashboard.js",
     "4-stat KPI row (Total Tasks, In Progress, Completed, High Priority) using StatCard components. "
     "Below that, a 2-column layout: left shows top 8 tasks sorted by priority using TaskCard components; "
     "right shows a compact completion ring panel (Progress Rings for done/active/critical %) and a "
     "Recent Activity feed. All data is fetched from GET /tasks and falls back to mock data if the "
     "backend is offline."),
    ("Projects.js",
     "Responsive auto-fill card grid using CSS grid(auto-fill, minmax(280px, 1fr)). Each ProjectCard "
     "shows the project name, creation date, 3 stat chips (total tasks, done, completion %), and a "
     "gradient progress bar. A 'New Project' button opens a Modal with a single input field. On submit, "
     "POST /projects is called and the new card is optimistically inserted at the top of the grid."),
    ("ProjectDetail.js",
     "3-column kanban board (Todo / In Progress / Done). Each column header shows a colored dot "
     "and task count. Tasks render as TaskCard components inside each column. An 'Add Task' panel "
     "slides open with three slider inputs for Impact (1-10), Effort (1-10), and Risk (1-10), plus "
     "an AI toggle that enables description generation via Gemini. Tasks are fetched from "
     "GET /projects/:id/tasks and PUT /tasks/prioritize is called after creation to re-rank globally."),
    ("Analytics.js",
     "4-stat row at top (same KPI cards). Main content: horizontal bar chart for task status "
     "distribution (Completed/In Progress/Backlog), a second horizontal bar chart for priority "
     "distribution (Critical/High/Medium/Low), and a weekly completions timeline (vertical bars). "
     "Sidebar shows 4 donut ProgressRing charts. Data from GET /analytics/kpis with mock fallback."),
    ("TeamFormation.js",
     "Left panel: skill chip-tag input (press Enter to add), team size slider (1-10), and a 'Suggest Team' "
     "button. Right panel: ranked candidate cards with initials avatar, name, email, color-highlighted "
     "skill chips, and a numeric match score (0-100). POST /teams/suggest sends the skills map and "
     "size limit; the backend AI scores each user and returns them sorted by score descending."),
]
for name, desc in pages:
    add_para(doc, desc, bold_prefix=name)

add_heading(doc, "4.6 Reusable UI Components", 2)
components = [
    ("TaskCard.js",      "Displays a single task. Left border is color-coded by priority level (red=Critical, orange=High, yellow=Medium, green=Low). Shows title, description (truncated), 3 metric pills (Impact/Effort/Risk), PriorityBadge, and StatusBadge."),
    ("StatCard.js",      "KPI metric card. Has a colored accent top bar, icon (SVG), large numeric value, label, and optional trend indicator (arrow up/down + percentage)."),
    ("PriorityBadge.js", "Maps a 0-100 score to a level: ≥75 Critical (red + pulsing dot), ≥50 High (orange), ≥25 Medium (yellow), else Low (green). Also contains StatusBadge for todo/in-progress/done/blocked."),
    ("ProgressRing.js",  "Pure SVG donut chart. Takes percent, size, strokeWidth, color, and label props. Uses SVG stroke-dasharray/dashoffset math to draw the arc. Has an animated transition on mount."),
    ("Modal.js",         "Overlay modal with backdrop blur. Click outside or press Escape to close. Accepts title, children (body), and footer props. Three size variants: sm, md, lg (controls max-width)."),
    ("EmptyState.js",    "Centered placeholder for empty lists. Accepts icon (SVG), title, description, and optional action (e.g., a button). Used in every list view to handle zero-data states gracefully."),
]
for name, desc in components:
    add_para(doc, desc, bold_prefix=name)

# ══════════════════════════════════════════════════════════════════════════════
# 5. BACKEND ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "5. Backend Architecture (Spring Boot)", 1, (79,70,229))

add_heading(doc, "5.1 Package Structure", 2)
for line in [
    "com.adapta/",
    "  AdaptaApplication.java      — @SpringBootApplication entry point",
    "  config/",
    "    FirebaseConfig.java       — Loads service account, initializes FirebaseApp",
    "    SecurityConfig.java       — Stateless security, injects FirebaseAuthFilter",
    "    CorsConfig.java           — Allows localhost:3000/3001, all methods",
    "  security/",
    "    FirebaseAuthFilter.java   — OncePerRequestFilter; validates Bearer JWT",
    "    FirebasePrincipal.java    — Value object: uid, email, name, photoURL",
    "  model/",
    "    User.java                 — @Document; embeds full profile",
    "    Project.java              — @Document; id, name, ownerId, timestamps",
    "    Task.java                 — @Document; scoring fields, priority, status",
    "    TaskStatus.java           — Enum: TODO, IN_PROGRESS, DONE, BLOCKED",
    "  repository/",
    "    UserRepository.java       — MongoRepository<User, String>",
    "    ProjectRepository.java    — findByOwnerIdOrderByCreatedAtDesc()",
    "    TaskRepository.java       — Derived queries by projectId, status, uid",
    "  dto/",
    "    ProfileRequest.java       — Onboarding form fields",
    "    TaskRequest.java          — Task creation payload",
    "    TeamSuggestRequest.java   — requiredSkills map + sizeLimit",
    "  service/",
    "    AuthService.java          — get/create user, save full profile",
    "    ProjectService.java       — CRUD for projects",
    "    TaskService.java          — Priority algorithm, AI description, reprioritize",
    "    TeamService.java          — Score users by skills/availability/experience",
    "    AnalyticsService.java     — Task count KPIs",
    "    AIService.java            — Gemini API wrapper (description + scoring)",
    "  controller/",
    "    AuthController.java       — GET /auth/me, POST /profile/complete",
    "    ProjectController.java    — GET/POST /projects, GET /projects/:id/tasks",
    "    TaskController.java       — GET/POST /tasks, PUT /tasks/prioritize",
    "    TeamController.java       — POST /teams/suggest",
    "    AnalyticsController.java  — GET /analytics/kpis",
]:
    add_code(doc, line)

add_heading(doc, "5.2 Request Lifecycle (End to End)", 2)
add_para(doc,
    "Here is the exact path of a POST /tasks request from browser to database and back:")
steps = [
    "1. React component calls: api.post('/tasks', { projectId, title, impact:8, effort:4, risk:3, generateDetails:true })",
    "2. Axios interceptor: attaches Authorization: Bearer eyJhbGci... to the request headers",
    "3. HTTP request arrives at Spring Boot on port 8080",
    "4. CorsFilter: validates Origin header matches allowed origins (localhost:3000)",
    "5. FirebaseAuthFilter.doFilterInternal(): extracts the Bearer token, calls FirebaseAuth.getInstance().verifyIdToken(token)",
    "6. Firebase Admin SDK validates the JWT signature against Google's public keys",
    "7. Filter creates a FirebasePrincipal(uid, email, name, photoURL) and sets it in SecurityContextHolder",
    "8. Request reaches TaskController.create(@AuthenticationPrincipal FirebasePrincipal p, @RequestBody TaskRequest req)",
    "9. TaskController calls taskService.createTask(p.getUid(), req)",
    "10. TaskService.createTask(): clamps impact/effort/risk to 1-10, computes urgencyBonus from deadline + project load",
    "11. Priority = clamp((8×8)+(3×4)-(4×2)+urgencyBonus, 0, 100) = clamp(64+12-8+bonus, 0, 100)",
    "12. Since generateDetails=true: calls aiService.generateTaskDescription(title, seedPrompt)",
    "13. AIService.callGemini(): HTTP POST to Google Gemini API, parses JSON response, returns description string",
    "14. Task document built with all fields; taskRepository.save(task) persists to MongoDB",
    "15. updateAssigneeWorkload(uid): re-computes activeTaskCount and currentWorkloadScore for the user",
    "16. Task returned as JSON → Axios → React state updated → UI re-renders with new task card",
]
for s in steps:
    add_bullet(doc, s)

add_heading(doc, "5.3 Security Design", 2)
add_para(doc,
    "Security is implemented as a stateless filter chain with no session cookies. Every request "
    "must carry a valid Firebase JWT. Here's the design:")
add_bullet(doc, "CSRF is disabled — not needed for stateless JWT APIs (no session/cookie to forge)")
add_bullet(doc, "Session creation policy = STATELESS — no HttpSession ever created")
add_bullet(doc, "FirebaseAuthFilter runs before Spring's UsernamePasswordAuthenticationFilter")
add_bullet(doc, "Token verification uses Google's cached public keys (no DB lookup needed)")
add_bullet(doc, "FirebasePrincipal is extracted by controllers via @AuthenticationPrincipal — no user lookup for every request")
add_bullet(doc, "If Firebase SDK is not initialized (missing service account in dev), the filter is skipped gracefully")
add_bullet(doc, "All CORS origins are configurable via application.yml — default: localhost:3000,3001")

# ══════════════════════════════════════════════════════════════════════════════
# 6. DATABASE DESIGN (MongoDB)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "6. Database Design (MongoDB)", 1, (79,70,229))

add_heading(doc, "6.1 Why MongoDB (Not SQL)?", 2)
comparisons = [
    ("Schema flexibility", "User profile fields evolve frequently (skills, certs added over time) — no ALTER TABLE needed", "Fixed schema requires migrations for any new field"),
    ("Document embedding", "Task and project data are always read together — embedding avoids JOINs", "Would need JOIN across users/projects/tasks tables"),
    ("Horizontal scaling", "Built-in sharding for high task volumes across organizations", "Vertical scaling only without significant config"),
    ("Developer speed", "Spring Data MongoDB's derived queries auto-generate queries from method names", "SQL requires explicit query strings or ORM mapping"),
]
add_table(doc,
    ["Concern", "MongoDB Approach", "SQL Approach"],
    comparisons
)

add_heading(doc, "6.2 Collections Schema", 2)

add_heading(doc, "Collection: users", 3)
for line in [
    "{",
    '  "_id": "firebase_uid_abc123",        // @Id — Firebase UID is the document ID',
    '  "email": "sarah@adapta.io",',
    '  "displayName": "Sarah Chen",',
    '  "photoURL": "https://...",',
    '  "hasProfile": true,',
    '',
    '  // ── Profile (for team formation) ──',
    '  "name": "Sarah Chen",',
    '  "title": "Senior Frontend Engineer",',
    '  "department": "Platform",',
    '  "location": "Bengaluru, India",',
    '  "timezone": "IST",',
    '  "yearsTotal": 5,',
    '  "skills": ["React", "TypeScript", "Java", "MongoDB"],',
    '  "certifications": ["AWS Solutions Architect"],',
    '  "availabilityHoursPerWeek": 40,',
    '  "preferences": ["remote", "async"],',
    '',
    '  // ── Computed workload (updated by TaskService) ──',
    '  "activeTaskCount": 2,',
    '  "currentWorkloadScore": 40.0,    // 0-100; 100 = fully loaded',
    '',
    '  "createdAt": ISODate("2025-12-01T10:00:00Z"),',
    '  "updatedAt": ISODate("2026-01-15T08:30:00Z")',
    "}",
]:
    add_code(doc, line)

add_heading(doc, "Collection: projects", 3)
for line in [
    "{",
    '  "_id": ObjectId("507f1f77bcf86cd799439011"),',
    '  "name": "Platform Redesign",',
    '  "ownerId": "firebase_uid_abc123",    // FK → users._id',
    '  "createdAt": ISODate("2026-01-01T00:00:00Z"),',
    '  "updatedAt": ISODate("2026-01-10T00:00:00Z")',
    "}",
]:
    add_code(doc, line)

add_heading(doc, "Collection: tasks", 3)
for line in [
    "{",
    '  "_id": ObjectId("507f191e810c19729de860ea"),',
    '  "projectId": "507f1f77bcf86cd799439011",  // FK → projects._id',
    '  "createdByUid": "firebase_uid_abc123",      // FK → users._id',
    '  "assignedToUid": null,',
    '',
    '  "title": "Redesign authentication flow",',
    '  "description": "Update login/signup UI...",   // possibly AI-generated',
    '',
    '  // ── Scoring inputs ──',
    '  "impact": 9,          // 1-10 (business value)',
    '  "effort": 4,          // 1-10 (engineering cost)',
    '  "risk": 3,            // 1-10 (risk of delay)',
    '',
    '  // ── Deadline ──',
    '  "deadline": ISODate("2026-03-14"),',
    '  "estimatedHours": 16,',
    '',
    '  // ── Computed (by TaskService) ──',
    '  "priority": 92.0,       // 0-100; sorted descending',
    '  "urgencyBonus": 30.0,   // from deadline proximity + project load',
    '',
    '  "status": "IN_PROGRESS",',
    '  "generateDetails": true,',
    '  "aiGenerated": true,',
    '',
    '  "createdAt": ISODate("2026-03-10T00:00:00Z"),',
    '  "updatedAt": ISODate("2026-03-11T00:00:00Z")',
    "}",
]:
    add_code(doc, line)

add_heading(doc, "6.3 Indexing Strategy", 2)
add_table(doc,
    ["Collection", "Field(s)", "Index Type", "Purpose"],
    [
        ["users",    "email",                  "Unique",     "Fast lookup by email; prevents duplicates"],
        ["tasks",    "projectId",              "Single",     "All queries filter by projectId first"],
        ["tasks",    "createdByUid",           "Single",     "Fetch all tasks for a user (Dashboard)"],
        ["tasks",    "assignedToUid",          "Single",     "Workload tracking for team formation"],
        ["tasks",    "projectId + status",     "Compound",   "countByProjectIdAndStatus for urgency bonus"],
        ["projects", "ownerId + createdAt",    "Compound",   "List projects for user, ordered newest first"],
    ]
)

add_heading(doc, "6.4 MongoDB Connection (Spring Data)", 2)
add_para(doc,
    "Spring Data MongoDB is configured via spring.data.mongodb.uri in application.yml. The "
    "MongoTemplate and MongoRepository beans are auto-configured by Spring Boot — no manual "
    "DataSource or SessionFactory setup needed. The Repository pattern provides derived query "
    "methods like findByProjectIdOrderByPriorityDesc(String projectId) which Spring Data "
    "translates directly to a MongoDB find({projectId:'...'}).sort({priority:-1}) query without "
    "any query string written by the developer.")

# ══════════════════════════════════════════════════════════════════════════════
# 7. AI FEATURES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "7. AI Features (Google Gemini)", 1, (79,70,229))

add_heading(doc, "7.1 Gemini API Integration", 2)
add_para(doc,
    "The AIService class wraps the Google Gemini 1.5 Flash REST API using Spring's RestTemplate "
    "(synchronous HTTP client). Gemini is called server-side only — the API key never reaches "
    "the browser. The endpoint used is:")
add_code(doc, "POST https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={API_KEY}")

add_para(doc, "Request JSON structure:")
for line in [
    '{',
    '  "contents": [{',
    '    "parts": [{ "text": "your prompt here" }]',
    '  }]',
    '}',
]:
    add_code(doc, line)

add_para(doc, "Response parsing:")
add_code(doc, "response → candidates[0] → content → parts[0] → text")

add_heading(doc, "7.2 Feature 1: AI Task Description Generation", 2)
add_para(doc,
    "When a user creates a task and toggles 'AI-Generated Description', the frontend sends "
    "generateDetails: true and an optional seedPrompt. The backend builds a prompt:")
add_code(doc, 'Prompt: "You are a software project manager. Write a concise 2-3 sentence task')
add_code(doc, ' description for: \\"Redesign authentication flow\\".')
add_code(doc, ' Additional context: Update login/signup UI to match new design.')
add_code(doc, ' Be specific, professional, and actionable. Plain text only."')
add_para(doc,
    "Gemini returns a naturally worded, context-aware description. This saves engineers time "
    "writing task details and ensures consistent formatting across the backlog.")

add_heading(doc, "7.3 Feature 2: AI Team Formation Scoring", 2)
add_para(doc,
    "When POST /teams/suggest is called, TeamService iterates all users with hasProfile=true "
    "and calls AIService.scoreTeamMember() for each. The Gemini prompt is:")
add_code(doc, 'Prompt: "Score this software engineer candidate 0-100 for a team where the required')
add_code(doc, ' skills are: React, Java. Candidate skills: React, TypeScript, Spring Boot.')
add_code(doc, ' Years of experience: 5. Available hours/week: 40. Active tasks: 2.')
add_code(doc, ' Consider semantic skill matches (Spring Boot counts for Java).')
add_code(doc, ' Reply with ONLY the integer score 0-100, nothing else."')
add_para(doc,
    "The key advantage over string comparison: Gemini understands that 'Spring Boot' implies "
    "Java knowledge, 'Next.js' implies React, 'Postgres' implies SQL — enabling nuanced matching "
    "that can't be done with simple toLowerCase().contains() checks.")
add_para(doc, "Fallback scoring formula (when no API key):")
add_code(doc, "skillMatch  = (matched skills / required skills count) × 100")
add_code(doc, "availScore  = (availableHours - activeTaskCount×8) / availableHours × 100")
add_code(doc, "expScore    = min(yearsTotal × 10, 100)")
add_code(doc, "SCORE       = skillMatch×0.60 + availScore×0.30 + expScore×0.10")

# ══════════════════════════════════════════════════════════════════════════════
# 8. PRIORITY ALGORITHM
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "8. Priority Algorithm (Full Detail)", 1, (79,70,229))

add_heading(doc, "8.1 The Formula", 2)
add_code(doc, "priority = clamp( (impact × 8) + (risk × 4) − (effort × 2) + urgencyBonus ,  0,  100 )")

add_heading(doc, "8.2 Factor Breakdown", 2)
add_table(doc,
    ["Factor", "Weight", "Range", "Rationale"],
    [
        ["impact",       "×8",  "1-10 → 8-80 pts",  "Business value is the most important driver of priority"],
        ["risk",         "×4",  "1-10 → 4-40 pts",  "Risky tasks need early attention to unblock others"],
        ["effort",       "×−2", "1-10 → -2 to -20", "Higher effort reduces priority (harder tasks take longer)"],
        ["urgencyBonus", "+",   "0-40 pts",          "Deadline proximity and project busyness add time pressure"],
    ]
)

add_heading(doc, "8.3 Urgency Bonus Calculation", 2)
add_table(doc,
    ["Condition", "Bonus Points", "Reasoning"],
    [
        ["Deadline is today or overdue",          "+30", "Maximum urgency — must be done immediately"],
        ["Deadline within 3 days",                "+30", "Critical deadline — top of the list"],
        ["Deadline within 4-7 days",              "+20", "High urgency — should be started today"],
        ["Deadline within 8-14 days",             "+10", "Moderate urgency — keep an eye on it"],
        ["No deadline set",                       " +0", "No time pressure — scored purely on impact/risk/effort"],
        ["Each in-progress task in project × 2",  "max +10", "Busy projects indicate higher overall urgency"],
    ]
)

add_heading(doc, "8.4 Worked Example", 2)
add_para(doc, "Task: 'Redesign authentication flow'")
add_code(doc, "  impact = 9  →  9  × 8 =  72")
add_code(doc, "  risk   = 3  →  3  × 4 =  12")
add_code(doc, "  effort = 4  →  4  × 2 =   8  (subtracted)")
add_code(doc, "  deadline in 2 days    = +30 (urgencyBonus)")
add_code(doc, "  project has 3 active  = +6  (workloadBonus, part of urgency)")
add_code(doc, "  ─────────────────────────────")
add_code(doc, "  raw    = 72 + 12 − 8 + 30 + 6 = 112")
add_code(doc, "  clamped to 100  →  priority = 100")

add_heading(doc, "8.5 Reprioritize Endpoint", 2)
add_para(doc,
    "PUT /tasks/prioritize recalculates the priority score for every task in the database. "
    "This is called automatically after each task creation. It fetches all tasks with "
    "taskRepository.findAll(), recomputes urgency (since days-to-deadline change daily), and "
    "batch-saves with taskRepository.saveAll(). This ensures the task list always reflects "
    "current time pressure even without any code changes.")

# ══════════════════════════════════════════════════════════════════════════════
# 9. USER PROFILE (TEAM FORMATION)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "9. User Profile — Complete Field Reference", 1, (79,70,229))
add_para(doc,
    "The User document is the most important model for team formation. It integrates "
    "Firebase identity data with rich professional profile fields collected during onboarding.")

add_table(doc,
    ["Field", "Type", "Source", "Used For"],
    [
        ["id (Firebase UID)",        "String",   "Firebase JWT sub claim", "Primary key, FK in tasks.createdByUid / assignedToUid"],
        ["email",                    "String",   "Firebase",               "Display in candidate cards; unique index"],
        ["displayName",              "String",   "Firebase (Google account)", "Fallback name before profile completion"],
        ["photoURL",                 "String",   "Firebase (Google avatar)", "Avatar display in team candidate cards"],
        ["hasProfile",               "boolean",  "Set after onboarding",   "Gate for onboarding redirect in App.js"],
        ["name",                     "String",   "Onboarding Step 1",      "Display name in team cards after profile"],
        ["title",                    "String",   "Onboarding Step 1",      "e.g. 'Senior Backend Engineer' — shown in candidate card"],
        ["department",               "String",   "Onboarding Step 1",      "Org context for team matching (same/diff dept)"],
        ["location",                 "String",   "Onboarding Step 1",      "Geographic context; timezone-aware scheduling"],
        ["timezone",                 "String",   "Onboarding Step 1",      "IST, UTC, PST etc. — for async team formation"],
        ["yearsTotal",               "int",      "Onboarding Step 2",      "Experience score: min(years × 10, 100) × 0.10"],
        ["skills",                   "String[]", "Onboarding Step 2",      "PRIMARY input for AI team scoring (semantic match)"],
        ["certifications",           "String[]", "Onboarding Step 2",      "Context for Gemini prompt (AWS, PMP, etc.)"],
        ["availabilityHoursPerWeek", "int",      "Onboarding Step 3",      "Availability score: (hours - activeTasks×8) / hours"],
        ["preferences",              "String[]", "Onboarding Step 3",      "Work style tags (remote, async, sprint-based)"],
        ["activeTaskCount",          "int",      "Computed by TaskService", "Re-calculated whenever a task is created/updated"],
        ["currentWorkloadScore",     "double",   "Computed by TaskService", "0-100: how loaded the person is right now (affects team score)"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 10. API ENDPOINTS REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "10. REST API — Complete Endpoint Reference", 1, (79,70,229))

add_table(doc,
    ["Method", "Endpoint", "Auth", "Request Body", "Response", "Notes"],
    [
        ["GET",  "/auth/me",              "✓", "—",                                   "{ uid, hasProfile, email, name }",       "Creates User doc if first login"],
        ["POST", "/profile/complete",     "✓", "ProfileRequest (all profile fields)", "{ success, hasProfile, uid }",           "Marks hasProfile=true"],
        ["GET",  "/projects",             "✓", "—",                                   "Project[] ordered by createdAt desc",    "Only current user's projects"],
        ["POST", "/projects",             "✓", "{ name }",                            "Project",                                "Sets ownerId = current user"],
        ["GET",  "/projects/{id}/tasks",  "✓", "—",                                   "Task[] ordered by priority desc",        "Scoped to a single project"],
        ["GET",  "/tasks",               "✓", "—",                                   "Task[] ordered by priority desc",        "All tasks created by current user"],
        ["POST", "/tasks",               "✓", "TaskRequest",                         "Task (with computed priority)",          "AI description if generateDetails=true"],
        ["PUT",  "/tasks/prioritize",    "✓", "—",                                   "{ status: 'reprioritized' }",            "Recalculates all tasks' priority scores"],
        ["GET",  "/analytics/kpis",      "✓", "—",                                   "{ total_tasks, done, in_progress, todo, completion_pct }", "Global task counts"],
        ["POST", "/teams/suggest",       "✓", "{ requiredSkills, sizeLimit }",       "ScoredUser[] sorted by score desc",      "Gemini semantic scoring with fallback"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 11. KEY DESIGN DECISIONS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "11. Design Decisions & Trade-offs", 1, (79,70,229))

decisions = [
    ("Firebase Auth over custom JWT",
     "Firebase handles token rotation, Google OAuth, password-less flows, and public key "
     "management automatically. Building equivalent infrastructure from scratch would take weeks "
     "and introduce security risks. Trade-off: vendor lock-in to Firebase."),
    ("MongoDB over PostgreSQL",
     "Task and user data have highly variable shapes (different skills, different deadline fields). "
     "MongoDB's document model allows schema-less evolution without migrations. The priority scoring "
     "is computationally trivial and doesn't require the transactional guarantees SQL provides. "
     "Trade-off: no ACID transactions across documents (acceptable for task management)."),
    ("Gemini 1.5 Flash (not GPT-4)",
     "Google Gemini is free-tier friendly, fast (Flash variant), and aligns with the Google "
     "ecosystem (Firebase, Cloud). Flash is the fastest Gemini model with sufficient quality for "
     "task descriptions and scoring prompts. Trade-off: slightly lower quality than GPT-4-turbo."),
    ("Mock data fallback on every page",
     "All pages display mock data when the backend is offline or returns an empty array. This "
     "means the UI is always functional regardless of backend state. This speeds up frontend "
     "development (no backend dependency) and provides a demo mode for stakeholders."),
    ("CSS Custom Properties (Design System) over Tailwind",
     "Vanilla CSS gives full control over the dark-mode design system with zero JavaScript overhead. "
     "CSS custom properties are native browser features — no build step, no class names to memorize. "
     "Trade-off: slightly more verbose CSS than utility-first frameworks."),
    ("Stateless Spring Security (no sessions)",
     "REST APIs should be stateless — every request is self-contained with a JWT. No server-side "
     "session storage means horizontal scaling works trivially (no sticky sessions or session "
     "replication needed). Trade-off: slightly more work to handle token revocation (not needed for MVP)."),
    ("Priority formula: impact×8, risk×4, effort×−2",
     "Weights were chosen to make impact the dominant driver (60%+ of the score), risk secondary "
     "(meaningful but not overwhelming), and effort a mild penalty (high effort = lower priority). "
     "This reflects a business-first prioritization philosophy. The formula is transparent, "
     "inspectable, and can be tuned by changing constants without changing architecture."),
]
for title, explanation in decisions:
    add_para(doc, explanation, bold_prefix=title)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 12. DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "12. Deployment Architecture", 1, (79,70,229))

add_heading(doc, "12.1 Environment Variables", 2)
add_table(doc,
    ["Variable", "Service", "Value/Description"],
    [
        ["REACT_APP_API_URL",           "Frontend", "http://localhost:8080 (dev) or backend URL (prod)"],
        ["REACT_APP_FIREBASE_API_KEY",  "Frontend", "Firebase project web API key"],
        ["REACT_APP_FIREBASE_AUTH_DOMAIN","Frontend","adapta-ba4b0.firebaseapp.com"],
        ["REACT_APP_FIREBASE_PROJECT_ID","Frontend", "adapta-ba4b0"],
        ["MONGODB_URI",                 "Backend",  "mongodb://localhost:27017/adapta or Atlas URI"],
        ["FIREBASE_SERVICE_ACCOUNT_PATH","Backend", "Path to firebase-service-account.json"],
        ["GEMINI_API_KEY",              "Backend",  "Google AI Studio API key"],
        ["cors.allowed-origins",        "Backend",  "http://localhost:3000 (dev) or production domain"],
    ]
)

add_heading(doc, "12.2 Running Locally", 2)
add_code(doc, "# Terminal 1: Start MongoDB")
add_code(doc, "mongod --dbpath C:/data/db")
add_code(doc, "")
add_code(doc, "# Terminal 2: Start Backend")
add_code(doc, "cd e:\\Adapta\\backend")
add_code(doc, "set GEMINI_API_KEY=your_key_here")
add_code(doc, "mvn spring-boot:run")
add_code(doc, "")
add_code(doc, "# Terminal 3: Start Frontend")
add_code(doc, "cd e:\\Adapta\\frontend")
add_code(doc, "npm start")
add_code(doc, "")
add_code(doc, "# Application: http://localhost:3000")
add_code(doc, "# API:         http://localhost:8080")
add_code(doc, "# MongoDB:     mongodb://localhost:27017/adapta")

add_heading(doc, "12.3 Production Deployment Plan", 2)
add_table(doc,
    ["Component", "Platform Option", "Notes"],
    [
        ["Frontend",  "Vercel / Netlify / Firebase Hosting", "Static build: npm run build, configure env vars in dashboard"],
        ["Backend",   "Railway / Render / Google Cloud Run",  "Dockerize: FROM eclipse-temurin:17-jre; ENTRYPOINT java -jar app.jar"],
        ["Database",  "MongoDB Atlas",                        "Free M0 cluster sufficient for MVP; automatic backups"],
        ["Auth",      "Firebase",                             "No deployment needed — fully managed SaaS"],
        ["AI",        "Google Gemini API",                    "No deployment needed — pay-per-use API"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 13. COMMON INTERVIEW QUESTIONS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "13. Common Interview Questions & Answers", 1, (79,70,229))

qas = [
    ("Q: How does Firebase authentication work end-to-end?",
     "A: Firebase Auth on the client handles OAuth (Google login) and returns a signed JWT ID Token. "
     "This token is attached to every API request as a Bearer header by the Axios interceptor. On "
     "the server, the Firebase Admin SDK verifies the token signature against Google's public keys "
     "(cached locally, refreshed periodically). On success, the decoded UID and email are stored "
     "as a FirebasePrincipal in Spring Security's SecurityContext. No password hashing, no session "
     "storage, no custom JWT signing — Firebase handles all of that."),
    ("Q: How is the priority score calculated?",
     "A: priority = clamp((impact×8) + (risk×4) - (effort×2) + urgencyBonus, 0, 100). "
     "Impact is weighted highest (×8) because business value is the primary driver. Risk adds "
     "urgency (×4) since risky tasks need early attention. Effort penalizes (-×2) since harder "
     "tasks naturally take longer. The urgencyBonus adds up to +30 for tasks with deadlines within "
     "3 days, plus up to +10 for busy projects (2 points per in-progress task in the same project)."),
    ("Q: How does the AI team formation scoring work?",
     "A: When POST /teams/suggest is called, TeamService fetches all profiled users. For each user, "
     "it calls Gemini with a prompt containing the required skills, candidate skills, years of experience, "
     "available hours/week, and active task count. Gemini returns a 0-100 integer score. The key "
     "advantage is semantic matching — Gemini knows that 'Spring Boot' implies Java, 'Next.js' implies "
     "React, so candidates aren't penalized for naming the specialization rather than the base technology. "
     "If no Gemini key is configured, it falls back to: skillMatch×0.60 + availability×0.30 + experience×0.10."),
    ("Q: Why use MongoDB instead of a relational database?",
     "A: Tasks have variable schemas — some have deadlines, some don't; skills arrays vary per user; "
     "certifications are added over time. MongoDB's document model handles this without schema migrations. "
     "Additionally, the primary query pattern is 'get all tasks for a project sorted by priority' — a "
     "simple single-collection find() with no JOINs needed. For the task management use case, MongoDB "
     "is faster to develop with and scales horizontally via sharding."),
    ("Q: How do you handle the backend being offline on the frontend?",
     "A: Every page has mock data fallback. API calls are wrapped in try/catch — if the backend returns "
     "an error or is unreachable, the page silently uses the MOCK_TASKS / MOCK_PROJECTS etc. arrays "
     "defined at the top of each page component. This means the UI is always interactive, which is "
     "critical for demos and early frontend development before the backend is ready."),
    ("Q: How does the Kanban board work?",
     "A: ProjectDetail.js fetches all tasks for a project from GET /projects/:id/tasks. Tasks are "
     "already sorted by priority descending. The frontend then splits them into 3 arrays using "
     ".filter(t => t.status === 'todo'), 'in-progress', and 'done'. Each column renders its "
     "array as TaskCard components. The columns themselves are CSS grid with 3 equal columns. "
     "Status changes are handled by calling PATCH /tasks/:id (if implemented) — for the MVP, "
     "status is set at creation time."),
    ("Q: How would you scale this to 10,000 users?",
     "A: Frontend: CDN-hosted static build on Vercel — already globally distributed. Backend: "
     "Containerize with Docker, deploy to Kubernetes on GKE — horizontal scaling. The stateless "
     "JWT auth means any pod can handle any request. MongoDB Atlas handles horizontal sharding "
     "automatically. The Gemini API is external and scales infinitely. The main bottleneck would "
     "be PUT /tasks/prioritize which does a full collection scan — that would need to be made "
     "incremental (only re-score tasks whose deadline changed or new tasks in the same project)."),
]
for q, a in qas:
    add_para(doc, "", bold_prefix=q)
    add_para(doc, a, indent=True)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = r"e:\Adapta\Adapta_Project_Documentation.docx"
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
